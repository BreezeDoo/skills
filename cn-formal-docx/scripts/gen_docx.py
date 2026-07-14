# -*- coding: utf-8 -*-
"""cn-formal-docx 生成脚本 (python-docx)

用 python-docx 生成中文规范文档(办法/规定/规程/大纲),复杂域(SEQ 题注、REF
交叉引用、PAGE 页码)在生成期用 lxml 直接构造,无需"占位 + 后处理正则替换"两段式。

用法:
  python gen_docx.py            # 用顶部 CONFIG 生成
  python gen_docx.py out.docx  # 覆盖输出名

工作流(单脚本,无后处理):
  1. 改顶部 CONFIG(章节 md、图目录、FIG_BY_CHAPTER、REF_BY_CHAPTER、OUT)。
  2. python gen_docx.py
  3. python scripts/validate/validate.py 输出.docx   # validate.py 已 vendor 进本 skill

只管排版格式,不管内容。内容(章节文字)由写作者填入 md。

正文字里图引用处不写占位:正文用 ref_para 生成 "见 图1、图2" 的 REF 域(指向题注
书签 _Ref_figN);题注用 caption_para 生成 SEQ 自动编号域。增删图后 Word 全选 F9
自动同步题注号与引用号(SEQ/REF 域均带 w:dirty="true",打开即提示更新)。
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, Mm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

# ============ 排版常量(按规范)============
CN = "仿宋_GB2312"            # 正文中文字体
CN_BOLD = "黑体"             # 题注中文字体
EN = "Times New Roman"       # 英文/数字字体
SIZE = 16                    # 正文三号 16pt
H1 = 16                      # 章标题三号
TITLE = 18                   # 大标题 18pt
CAPTION = 12                 # 题注/表格小四 12pt
FOOTER_SZ = 10.5             # 页脚小五
LINE = 360                   # 1.5 倍行距 (240*1.5)

# 图尺寸限制 (px, 1px=9525 EMU @96dpi)
LANDSCAPE_USABLE_W = 920     # 横排可用宽
MAX_FIG_H = 520              # 横排图高限(保证图+题注同页)
INLINE_MAX_W = 530           # 内联小图最大宽
INLINE_MAX_H = 480           # 内联小图最大高

EMU_PER_PX = 9525

# ============ 输入配置(填入实际内容)============
CONFIG = {
    "MD": "仿真资源管理办法.md",
    "FIG_DIR": "仿真资源管理架构图-v2",
    "OUT": "仿真资源管理办法_py.docx",
    # 章节 -> 图列表(顺序即 SEQ 编号, 全局连续)
    # layout: "landscape"(大图横排分节符) | "inline"(小图正文内联) | "auto"(按宽高比)
    "FIG_BY_CHAPTER": {
        "第二章 体系结构": [
            {"file": "sim-resource-management-overview-v2.drawio.png", "name": "仿真资源管理总览", "layout": "landscape"},
            {"file": "sim-resource-taxonomy-assembly-v2.drawio.png", "name": "资源分类与工程装配", "layout": "landscape"},
        ],
        "第五章 更新管理": [
            {"file": "sim-resource-version-traceability-v2.drawio.png", "name": "版本追溯关系", "layout": "landscape"},
        ],
        "第九章 配置审核与追溯": [
            {"file": "sim-resource-lifecycle-governance-v2.drawio.png", "name": "资源生命周期与治理", "layout": "landscape"},
        ],
    },
    # 章节 -> 该章末尾图引用句(指向 SEQ 编号)
    "REF_BY_CHAPTER": {
        "第二章 体系结构": {"prefix": "仿真资源体系结构与装配关系见", "refs": [1, 2], "suffix": "。"},
        "第五章 更新管理": {"prefix": "资源更新与版本演进关系见", "refs": [3], "suffix": "。"},
        "第九章 配置审核与追溯": {"prefix": "版本追溯与配置审核流程见", "refs": [4], "suffix": "。"},
    },
}

# ============ 排版工具函数 ============
# OOXML 元素顺序约束: tblPr/tcPr 子元素必须按 schema 顺序排列, 否则 validate 失败。
# insert_ordered 按目标顺序把元素插到正确位置(在第一个序号更大的元素之前), 兼容已有子元素。
TBLPR_ORDER = ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
               "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
               "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription", "tblPrChange"]
TCPR_ORDER = ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
              "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "tcPrChange"]


def _local(tag):
    return tag.split("}")[-1]


def insert_ordered(parent, elem, order):
    """把 elem 按顺序表插入 parent(在第一个序号更大的已有元素之前, 否则末尾)。"""
    try:
        idx = order.index(_local(elem.tag))
    except ValueError:
        parent.append(elem)
        return
    for child in parent:
        try:
            cidx = order.index(_local(child.tag))
        except ValueError:
            continue
        if cidx > idx:
            child.addprevious(elem)
            return
    parent.append(elem)


# 字体规格: (中文, 英文, 字号pt, 加粗)
BODY_FONT = (CN, EN, SIZE, False)
TITLE_FONT = (CN, EN, TITLE, True)
H1_FONT = (CN, EN, H1, True)
CAP_FONT = (CN_BOLD, EN, CAPTION, True)   # 题注黑体小四加粗
TABLE_FONT = (CN, EN, CAPTION, False)
FOOTER_FONT = (CN, EN, FOOTER_SZ, False)


def apply_font(run, spec):
    """中英分体: ascii/hAnsi=英文, eastAsia=中文。必须用对象形式, 不能只设字符串。"""
    cn_font, en_font, size_pt, bold = spec
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = en_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)


def set_spacing(paragraph, line=LINE, before=None, after=None):
    """显式设行距 line + lineRule=auto(坑: 只设 line 部分 Word 不识别 1.5 倍)。"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line / 240.0   # 1.5 倍; python-docx 转 line+lineRule=auto
    if before is not None:
        pf.space_before = Pt(before / 20.0) if before > 100 else Pt(before)
    if after is not None:
        pf.space_after = Pt(after / 20.0) if after > 100 else Pt(after)


def add_field(paragraph, instr, cached_value, font_spec, bookmark=None):
    """在段落末尾追加复杂域(fldChar begin/instrText/separate/缓存值/end)。

    - instr: 域指令, 如 'SEQ Figure \\* ARABIC' / 'REF _Ref_fig1 \\h' / 'PAGE'
    - cached_value: 域结果缓存(打开时显示, F9 后刷新); SEQ/REF 用序号, PAGE 用 '1'
    - bookmark: (id, name) 仅 SEQ 域需要, 包裹缓存值 run, 供 REF 定位
    - begin 带 w:dirty="true": Word 打开主动提示更新域(借鉴 inject_toc_field.py)
    """
    def add_field_run():
        r = paragraph.add_run()
        apply_font(r, font_spec)
        return r

    # begin (dirty)
    r = add_field_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fld.set(qn("w:dirty"), "true")
    r._r.append(fld)
    # instrText (前后留空格, xml:space=preserve)
    r = add_field_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " " + instr.strip() + " "
    r._r.append(it)
    # separate
    r = add_field_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)
    # bookmark start (SEQ 域包裹缓存值)
    if bookmark:
        bid, bname = bookmark
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), str(bid))
        bs.set(qn("w:name"), bname)
        paragraph._p.append(bs)
    # 缓存值 run
    r = add_field_run()
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = str(cached_value)
    r._r.append(t)
    # bookmark end
    if bookmark:
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), str(bid))
        paragraph._p.append(be)
    # end
    r = add_field_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


# ============ 段落工厂 ============
def title_para(doc, text):
    """大标题: 居中加粗, 不进 Heading(否则与章混在导航树)。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=18)
    r = p.add_run(text)
    apply_font(r, TITLE_FONT)
    return p


def h1_para(doc, text):
    """章标题: Heading1 + outlineLevel:0(导航窗格可见)。
    用 docx 内置 Heading 1 样式并在 styles 覆盖其字体(见 configure_styles)。"""
    p = doc.add_paragraph(style="Heading 1")
    set_spacing(p, before=18, after=9)
    r = p.add_run(text)
    apply_font(r, H1_FONT)
    # 确保 outlineLevel:0
    pPr = p._p.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        pPr.append(ol)
    ol.set(qn("w:val"), "0")
    return p


def body_para(doc, text):
    """正文段: 仿宋三号, 1.5 倍, 首行缩进 2 字符(三号=32pt=640twips)。"""
    p = doc.add_paragraph()
    set_spacing(p)
    p.paragraph_format.first_line_indent = Pt(32)
    r = p.add_run(text)
    apply_font(r, BODY_FONT)
    return p


def ref_para(doc, prefix, refs, suffix):
    """图引用句: "prefix 图REF1、图REF2 suffix", REF 域指向题注书签 _Ref_figN。"""
    p = doc.add_paragraph()
    set_spacing(p)
    p.paragraph_format.first_line_indent = Pt(32)
    r = p.add_run(prefix)
    apply_font(r, BODY_FONT)
    for i, n in enumerate(refs):
        if i > 0:
            r = p.add_run("、")
            apply_font(r, BODY_FONT)
        r = p.add_run("图")
        apply_font(r, BODY_FONT)
        add_field(p, "REF _Ref_fig%d \\h" % n, str(n), BODY_FONT)
    r = p.add_run(suffix)
    apply_font(r, BODY_FONT)
    return p


def caption_para(doc, name, seq):
    """题注: 居中黑体小四加粗 "图 SEQ  name", SEQ 域缓存值外包书签 _Ref_figN。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=6, after=12)
    p.paragraph_format.keep_with_next = True   # 图与题注同页
    r = p.add_run("图 ")
    apply_font(r, CAP_FONT)
    add_field(p, "SEQ Figure \\* ARABIC", str(seq), CAP_FONT,
              bookmark=(1000 + seq, "_Ref_fig%d" % seq))
    r = p.add_run("  " + name)
    apply_font(r, CAP_FONT)
    return p


# ============ 图 ============
def img_size(file, fig_dir):
    im = Image.open(os.path.join(fig_dir, file))
    w, h = im.size
    return w, h, w / h


def fig_landscape_para(doc, file, fig_dir, is_page_break):
    """横排大图: 按高反算宽(高限520, 宽不超920), 保证图+题注同页。"""
    w0, h0, ratio = img_size(file, fig_dir)
    h = MAX_FIG_H
    w = round(h * ratio)
    if w > LANDSCAPE_USABLE_W:
        w = LANDSCAPE_USABLE_W
        h = round(w / ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=0)
    p.paragraph_format.keep_with_next = True
    if is_page_break:
        p.paragraph_format.page_break_before = True
    r = p.add_run()
    r.add_picture(os.path.join(fig_dir, file),
                  width=Emu(w * EMU_PER_PX), height=Emu(h * EMU_PER_PX))
    return p


def fig_inline_para(doc, file, fig_dir):
    """内联小图: 按竖排版心宽缩放(宽限530, 高限480), 不分节, 留在正文流。"""
    w0, h0, ratio = img_size(file, fig_dir)
    w = INLINE_MAX_W
    h = round(w / ratio)
    if h > INLINE_MAX_H:
        h = INLINE_MAX_H
        w = round(h * ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=6, after=0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(os.path.join(fig_dir, file),
                  width=Emu(w * EMU_PER_PX), height=Emu(h * EMU_PER_PX))
    return p


def resolve_layout(fig, fig_dir):
    layout = fig.get("layout", "auto")
    if layout in ("landscape", "inline"):
        return layout
    _, _, ratio = img_size(fig["file"], fig_dir)
    return "landscape" if ratio > 1.3 else "inline"


# ============ 表格 ============
def add_md_table(doc, rows):
    """markdown 表格 -> 三线表风格: 全边框999999, 表头底色E8EEF5加粗, 小四。"""
    ncols = len(rows[0])
    total_w = 9360
    col_w = total_w // ncols
    widths = [col_w] * (ncols - 1) + [total_w - col_w * (ncols - 1)]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.autofit = False

    tblPr = table._tbl.tblPr
    # tblW (不在则创建, 按序插入)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        insert_ordered(tblPr, tblW, TBLPR_ORDER)
    tblW.set(qn("w:w"), str(total_w))
    tblW.set(qn("w:type"), "dxa")
    # 表格居中: Transitional schema CT_TblPrBase 用 w:jc (非 tblJc, tblJc 会被 schema 拒)
    # 不用 table.alignment setter: 它写 tblJc 且位置错
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        insert_ordered(tblPr, jc, TBLPR_ORDER)
    jc.set(qn("w:val"), "center")
    # tblBorders
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        insert_ordered(tblPr, borders, TBLPR_ORDER)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(qn("w:%s" % edge))
        if b is None:
            b = OxmlElement("w:%s" % edge)
            borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "999999")

    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            cell = table.cell(ri, ci)
            cell.width = Twips(widths[ci])   # 设 tcW
            tcPr = cell._tc.get_or_add_tcPr()
            # 表头底色 shd (按序插入, 保证在 tcMar 之前)
            if ri == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "E8EEF5")
                insert_ordered(tcPr, shd, TCPR_ORDER)
            # CT_TcMar 顺序: top, [start], left, bottom, [end], right (left 必须在 bottom 之前)
            tcMar = OxmlElement("w:tcMar")
            for edge, val in (("top", "30"), ("left", "100"), ("bottom", "30"), ("right", "100")):
                m = OxmlElement("w:%s" % edge)
                m.set(qn("w:w"), val)
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            insert_ordered(tcPr, tcMar, TCPR_ORDER)
            # 文字
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(cell_text)
            apply_font(r, (CN, EN, CAPTION, ri == 0))
    doc.add_paragraph()   # 表后留空
    return table


# ============ 分节符与页脚 ============
def configure_section(section, orientation):
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def set_footer_pagenum(section):
    """页脚居中页码(无页眉)。断开继承, 每个 section 独立设 PAGE 域。"""
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r1 = p.add_run("— ")
    apply_font(r1, FOOTER_FONT)
    add_field(p, "PAGE", "1", FOOTER_FONT)
    r3 = p.add_run(" —")
    apply_font(r3, FOOTER_FONT)


# ============ 默认样式 ============
def configure_styles(doc):
    """覆盖 Normal 与 Heading 1 样式: 中英分体三号、Heading1 三号加粗 outlineLevel:0。"""
    normal = doc.styles["Normal"]
    normal.font.size = Pt(SIZE)
    normal.font.name = EN
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), EN)
    rfonts.set(qn("w:hAnsi"), EN)
    rfonts.set(qn("w:eastAsia"), CN)
    # 默认 1.5 倍行距
    pPr = normal.element.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), str(LINE))
    sp.set(qn("w:lineRule"), "auto")

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(H1)
    h1.font.bold = True
    h1.font.name = EN
    rpr = h1.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), EN)
    rfonts.set(qn("w:hAnsi"), EN)
    rfonts.set(qn("w:eastAsia"), CN)
    pPr = h1.element.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), str(LINE))
    sp.set(qn("w:lineRule"), "auto")
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        pPr.append(ol)
    ol.set(qn("w:val"), "0")


# ============ 状态: 当前节方向 ============
class Builder:
    def __init__(self, doc, fig_dir, fig_by_chapter, ref_by_chapter):
        self.doc = doc
        self.fig_dir = fig_dir
        self.fig_by_chapter = fig_by_chapter
        self.ref_by_chapter = ref_by_chapter
        self.orientation = "portrait"
        self.seq = 0

    def ensure_portrait(self):
        if self.orientation != "portrait":
            sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(sec, "portrait")
            set_footer_pagenum(sec)
            self.orientation = "portrait"

    def ensure_landscape(self):
        if self.orientation != "landscape":
            sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(sec, "landscape")
            set_footer_pagenum(sec)
            self.orientation = "landscape"

    def flush_chapter(self, chapter):
        """进新章前: 先插引用句, 再插图(随章节, 非堆附录)。"""
        ref = self.ref_by_chapter.get(chapter)
        if ref:
            self.ensure_portrait()
            ref_para(self.doc, ref["prefix"], ref["refs"], ref["suffix"])
        figs = self.fig_by_chapter.get(chapter)
        if not figs:
            return
        pending_landscape = []
        for fig in figs:
            self.seq += 1
            layout = resolve_layout(fig, self.fig_dir)
            if layout == "inline":
                if pending_landscape:
                    self._emit_landscape_group(pending_landscape)
                    pending_landscape = []
                self.ensure_portrait()
                fig_inline_para(self.doc, fig["file"], self.fig_dir)
                caption_para(self.doc, fig["name"], self.seq)
            else:
                pending_landscape.append((fig, self.seq))
        if pending_landscape:
            self._emit_landscape_group(pending_landscape)

    def _emit_landscape_group(self, figs):
        """同章多横排图合并一个横排 section, 图间 pageBreakBefore(避免空页)。"""
        self.ensure_landscape()
        for i, (fig, seq) in enumerate(figs):
            fig_landscape_para(self.doc, fig["file"], self.fig_dir, is_page_break=(i > 0))
            caption_para(self.doc, fig["name"], seq)


# ============ 解析 markdown 正文 ============
def parse_md(md_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")
    return lines


def build(doc, lines, builder):
    pending_chapter = None
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # 表格
        if stripped.startswith("|"):
            buf = []
            while i < n and lines[i].strip().startswith("|"):
                buf.append(lines[i].strip())
                i += 1
            rows_all = [r for r in buf if not re.match(r"^\|[-:\s|]+$", r)]
            parsed = []
            for r in rows_all:
                cells = [c.strip() for c in r.split("|")]
                if cells and cells[0] == "":
                    cells = cells[1:]
                if cells and cells[-1] == "":
                    cells = cells[:-1]
                parsed.append(cells)
            builder.ensure_portrait()
            add_md_table(doc, parsed)
            continue
        # 大标题
        if line.startswith("# ") and not line.startswith("## "):
            builder.ensure_portrait()
            title_para(doc, line[2:].strip())
            i += 1
            continue
        # 章标题
        if line.startswith("## "):
            if pending_chapter:
                builder.flush_chapter(pending_chapter)
            ch_text = line[3:].strip()
            builder.ensure_portrait()
            h1_para(doc, ch_text)
            pending_chapter = ch_text
            i += 1
            continue
        # 普通段落(连续非空行合并)
        buf = []
        while i < n and lines[i].strip() and not lines[i].startswith("#") and not lines[i].strip().startswith("|"):
            buf.append(lines[i].strip())
            i += 1
        if buf:
            builder.ensure_portrait()
            body_para(doc, "".join(buf))
    if pending_chapter:
        builder.flush_chapter(pending_chapter)


# ============ main ============
def main():
    cfg = CONFIG
    if len(sys.argv) > 1:
        cfg = dict(CONFIG)
        cfg["OUT"] = sys.argv[1]
    out = os.path.abspath(cfg["OUT"])
    md_path = cfg["MD"]
    fig_dir = cfg["FIG_DIR"]

    # 切到 md 所在目录(图目录相对它)
    base = os.path.dirname(os.path.abspath(md_path)) or "."
    os.chdir(base)
    fig_dir = cfg["FIG_DIR"]

    doc = Document()
    configure_styles(doc)

    sec0 = doc.sections[0]
    configure_section(sec0, "portrait")
    set_footer_pagenum(sec0)

    # 首页空白
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    builder = Builder(doc, fig_dir, cfg["FIG_BY_CHAPTER"], cfg["REF_BY_CHAPTER"])
    lines = parse_md(cfg["MD"])
    build(doc, lines, builder)

    # python-docx 默认 settings.xml 的 zoom 缺 percent 属性, 补上(validate 要求)
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    doc.save(out)
    print("SAVED:", out)
    print("sections:", len(doc.sections))


if __name__ == "__main__":
    main()
